import io
import re
import json
import base64
from pathlib import Path
from typing import List, Dict
from datetime import datetime
import threading
import sys
import time
import os
import tempfile
import requests
from urllib.parse import urlparse
import uuid
from flask import Blueprint, request, jsonify, send_file
from flask_cors import CORS
import logging
from datetime import datetime, timedelta
from io import BytesIO
from fpdf import FPDF

from openai import AzureOpenAI
import httpx
import PyPDF2
import pandas as pd
import docx
from PIL import Image
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions, ContentSettings
from azure.cosmos import CosmosClient, PartitionKey, exceptions

# Configure logging
logger = logging.getLogger(__name__)

# Initialize Cosmos DB client
cosmos_client = CosmosClient(
    "https://barclaysdb.documents.azure.com:443/",
    "ercuc7wFNt4RPsxcx2QTzKP8AhKUDnjJrt0mpZrW2Yi2IQGAa7wDrEbhKRHsurGu0P7GuGny4IZRACDbtecfNQ=="
)
database = cosmos_client.get_database_client("RequirementsDB")
cosmos_db = database.get_container_client("StandardDocs")

# Create blueprint
srs_brd_bp = Blueprint('srs_brd', __name__, url_prefix='/srs_brd')
CORS(srs_brd_bp)  # Enable CORS for this blueprint

# Azure OpenAI Configuration
AZURE_ENDPOINT = "https://suraj-m9lgdbv9-eastus2.cognitiveservices.azure.com/"
AZURE_API_KEY = "75PVa3SAy9S2ZR590gZesyTNDMZtb3Oa5EdHlRbqWeQ89bmoOGl4JQQJ99BDACHYHv6XJ3w3AAAAACOGUP8d"
AZURE_API_VERSION = "2024-02-15-preview"
DEPLOYMENT_NAME = "gpt-4o"

# Azure Document Intelligence Configuration
FORM_ENDPOINT = "https://barclaysform.cognitiveservices.azure.com/"
FORM_KEY = "63spGg0VYFV0kWZB3nmsFDp8yEbi40zmEnCvIl6D8Seih4YyLsp9JQQJ99BDACYeBjFXJ3w3AAALACOGh5hu"

# Azure Cosmos DB Configuration
COSMOS_ENDPOINT = "https://barclaysdb.documents.azure.com:443/"
COSMOS_KEY = "ercuc7wFNt4RPsxcx2QTzKP8AhKUDnjJrt0mpZrW2Yi2IQGAa7wDrEbhKRHsurGu0P7GuGny4IZRACDbtecfNQ=="
COSMOS_DATABASE = "RequirementsDB"
COSMOS_CONTAINER = "StandardDocs"

# Initialize Cosmos DB client with error handling
try:
    cosmos_client = CosmosClient(COSMOS_ENDPOINT, COSMOS_KEY)
    database = cosmos_client.get_database_client(COSMOS_DATABASE)
    container = database.get_container_client(COSMOS_CONTAINER)
    print("Successfully connected to Cosmos DB")
except Exception as e:
    print(f"Error connecting to Cosmos DB: {e}")
    cosmos_client = None
    database = None
    container = None

def spinner(stop_event):
    """A simple console spinner to indicate processing."""
    spinner_chars = ['|', '/', '-', '\\']
    idx = 0
    while not stop_event.is_set():
        sys.stdout.write(f"\rGenerating document... {spinner_chars[idx % len(spinner_chars)]}")
        sys.stdout.flush()
        idx += 1
        time.sleep(0.1)
    sys.stdout.write("\rGeneration complete!             \n")

class MultiMediaProcessor:
    def __init__(self):
        """
        Initialize the MultiMediaProcessor with Azure OpenAI and Document Intelligence.
        """
        self.client = AzureOpenAI(
            azure_endpoint=AZURE_ENDPOINT,
            api_key=AZURE_API_KEY,
            api_version=AZURE_API_VERSION
        )
        self.form_recognizer_client = DocumentAnalysisClient(
            endpoint=FORM_ENDPOINT,
            credential=AzureKeyCredential(FORM_KEY)
        )
        self.original_source_content = ""
        self.original_source_names = []
        self.document_name = ""
        self.business_domain = None
        self.generated_sections = {}
        self.section_evaluations = {}

    def _encode_image(self, image_path: str) -> str:
        """Encode an image to base64."""
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def extract_text(self, file_path: str) -> str:
        """Extract text from various file types using Azure Document Intelligence when possible."""
        path = Path(file_path)
        try:
            # First try to use Azure Document Intelligence for supported formats
            if path.suffix.lower() in ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.docx']:
                with open(path, "rb") as document:
                    poller = self.form_recognizer_client.begin_analyze_document(
                        "prebuilt-document", document
                    )
                    result = poller.result()
                
                text_content = []
                for paragraph in result.paragraphs:
                    text_content.append(paragraph.content)
                
                # For structured documents, also extract any tables
                if hasattr(result, 'tables'):
                    for table in result.tables:
                        table_content = []
                        for cell in table.cells:
                            table_content.append(cell.content)
                        text_content.append("\n".join(table_content))
                
                return "\n".join(text_content)
            
            # Fallback to traditional methods for other formats
            elif path.suffix.lower() in ['.xls', '.xlsx']:
                return self._extract_excel_text(path)
            elif path.suffix.lower() == '.txt':
                return self._extract_txt_text(path)
            else:
                print(f"Unsupported file type: {path.suffix}")
                return ""
                
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return ""

    def _process_image(self, path: Path) -> str:
        """Process and extract text from an image using Azure Document Intelligence."""
        try:
            with open(path, "rb") as document:
                poller = self.form_recognizer_client.begin_analyze_document(
                    "prebuilt-document", document
                )
                result = poller.result()

            # Extract text from paragraphs
            text_content = []
            for paragraph in result.paragraphs:
                text_content.append(paragraph.content)

            return "\n".join(text_content)
        except Exception as e:
            print(f"Error processing image {path}: {e}")
            return ""

    def _extract_pdf_text(self, path: Path) -> str:
        """Extract text from a PDF."""
        with path.open('rb') as file:
            reader = PyPDF2.PdfReader(file)
            return '\n'.join(page.extract_text() for page in reader.pages)

    def _extract_docx_text(self, path: Path) -> str:
        """Extract text from a DOCX file."""
        doc = docx.Document(path)
        return '\n'.join(paragraph.text for paragraph in doc.paragraphs)

    def _extract_excel_text(self, path: Path) -> str:
        """Extract text from an Excel file."""
        df = pd.read_excel(path)
        return df.to_string()

    def _extract_txt_text(self, path: Path) -> str:
        """Extract text from a plain text file."""
        return path.read_text(encoding='utf-8')

    def process_multiple_documents(self,
                                   document_type: str,
                                   source_files: List[str],
                                   business_domain: str = None) -> Dict[str, str]:
        """
        Combine the content of all source files and generate one document.
        """
        combined_content = ""
        source_names = []

        for file_path in source_files:
            try:
                extracted_text = self.extract_text(file_path)
                file_name = Path(file_path).name
                combined_content += f"\n--- Content from {file_name} ---\n{extracted_text}\n"
                source_names.append(file_name)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")

        # Store original source content and related details for later regeneration
        self.original_source_content = combined_content
        self.original_source_names = source_names

        # Generate document name based on date and type
        current_date = datetime.datetime.now().strftime("%Y%m%d")
        self.document_name = f"{document_type}_{current_date}"
        self.business_domain = business_domain

        # Generate document content in sections
        generated_doc = self._generate_document_in_sections(
            document_type,
            combined_content,
            self.document_name,
            source_names,
            business_domain
        )

        return {self.document_name: generated_doc}

    def _generate_document_in_sections(self,
                                         document_type: str,
                                         source_content: str,
                                         document_name: str,
                                         source_names: List[str],
                                         business_domain: str = None) -> str:
        """Generate concise documents: BRD (8-10 pages) and SRS (2-3 pages)."""
        if document_type.upper() == 'BRD':
            sections = [
                "Executive Summary",
                "Project Objectives",
                "Project Scope",
                "Business Requirements",
                "Key Stakeholders",
                "Project Constraints",
                "Cost-Benefit Analysis",
                "MoSCoW Prioritization",
                "References"
            ]
            prompt_template = (
                "Generate a CONCISE section titled '{section}' for the Business Requirements Document. "
                "The entire BRD should be 8-10 pages, so keep this section proportionally sized. "
                "Use bullet points for clarity. Focus on key information without verbose explanations.\n\n"
                "Business Domain Context: {business_context}\n\n"
                "Source Content:\n{source_content}\n\n"
                "Internal Sources: {internal_sources}\n"
            )
        elif document_type.upper() == 'SRS':
            sections = [
                "Introduction",
                "System Overview",
                "Functional Requirements",
                "Non-Functional Requirements",
                "Interface Requirements",
                "Design Constraints"
            ]
            prompt_template = (
                "Generate a VERY CONCISE section titled '{section}' for the Software Requirements Specification. "
                "The entire SRS must be 2-3 pages, so keep this section brief and focused. "
                "Use bullet points and requirement IDs (REQ-XXX) for clear organization.\n\n"
                "Business Domain Context: {business_context}\n\n"
                "Source Content:\n{source_content}\n\n"
                "Internal Sources: {internal_sources}\n"
            )
        else:
            raise ValueError("Document type must be 'BRD' or 'SRS'")

        business_context = business_domain if business_domain else "General"
        internal_sources = ", ".join(source_names)

        sections_content = []
        for section in sections:
            prompt = prompt_template.format(
                section=section,
                document_name=document_name,
                business_context=business_context,
                source_content=source_content,
                internal_sources=internal_sources
            )
            try:
                # Calculate token limits based on document type and target page length
                if document_type.upper() == 'BRD':
                    # For 8-10 pages, allocate ~1000 tokens per page, distributed across sections
                    if section in ["Business Requirements", "Project Scope"]:
                        max_tokens = 1500  # More detailed sections
                    elif section == "Executive Summary":
                        max_tokens = 500   # Brief overview
                    else:
                        max_tokens = 1000  # Standard sections
                else:  # SRS
                    # For 2-3 pages, allocate ~500 tokens per page, distributed across sections
                    if section == "Functional Requirements":
                        max_tokens = 400   # Core requirements
                    elif section in ["Introduction", "System Overview"]:
                        max_tokens = 200   # Brief sections
                    else:
                        max_tokens = 300   # Standard sections

                response = self.client.chat.completions.create(
                    model=DEPLOYMENT_NAME,
                    messages=[
                        {"role": "system", "content": "You are an expert requirements writer. Generate clear, concise content using bullet points. Keep sections brief but informative. DO NOT include the section title in your response, as it will be added automatically."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=max_tokens,
                    temperature=0.7,
                    top_p=0.9
                )
                
                content = response.choices[0].message.content
                
                # Format bullet points consistently
                formatted_lines = []
                for line in content.split('\n'):
                    line = line.strip()
                    if line:
                        # Skip lines that look like section headers (they'll be added automatically)
                        if line.startswith('#') or line.startswith('##') or line.startswith('###'):
                            continue
                        # Convert various bullet point styles to a consistent format
                        if line.startswith('- ') or line.startswith('* '):
                            line = '• ' + line[2:]
                        elif line.startswith('-') or line.startswith('*'):
                            line = '• ' + line[1:]
                        formatted_lines.append(line)
                
                formatted_content = '\n'.join(formatted_lines)
                
                # Add section header
                section_text = f"## {section}\n\n{formatted_content}\n"
                sections_content.append(section_text)
                self.generated_sections[section] = formatted_content
                
            except Exception as e:
                print(f"Error generating section {section}: {e}")
                sections_content.append(f"## {section}\n\nError generating content for this section.\n")
        
        generated_doc = "\n".join(sections_content)
        return generated_doc

    def evaluate_section_completeness(self, section_name: str, section_content: str) -> Dict:
        """Evaluate if a section has all necessary information."""
        try:
            # Create a more specific prompt that considers the section type and content
            prompt = f"""Analyze this {section_name} section and identify any missing or incomplete information.
            Consider the following aspects:
            1. Content completeness for this specific section type
            2. Required elements for {section_name}
            3. Quality and depth of information
            4. Consistency with other sections
            
            Respond in valid JSON format with exactly these fields:
            {{
                "completeness_score": <number between 0-100, must be unique for each section>,
                "missing_elements": [<list of strings describing missing items>],
                "suggestions": [<list of strings with improvement suggestions>]
            }}

            Section Content:
            {section_content}
            """
            
            response = self.client.chat.completions.create(
                model=DEPLOYMENT_NAME,
                messages=[
                    {"role": "system", "content": """You are an expert requirements document analyst. 
                     Always respond with valid JSON containing exactly these fields:
                     {{
                         "completeness_score": <number between 0-100, must be unique for each section>,
                         "missing_elements": [<list of strings>],
                         "suggestions": [<list of strings>]
                     }}
                     
                     IMPORTANT: The completeness_score must be different for each section based on its actual content quality and completeness.
                     Do not use the same score for multiple sections."""},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7  # Increased temperature for more varied responses
            )
            
            try:
                # Get the response text and clean it to ensure valid JSON
                response_text = response.choices[0].message.content.strip()
                response_text = re.sub(r'^```json\s*|\s*```$', '', response_text)
                response_text = re.sub(r'^```\s*|\s*```$', '', response_text)
                
                # Parse the JSON response
                evaluation = json.loads(response_text)
                
                # Validate the required fields
                required_fields = ["completeness_score", "missing_elements", "suggestions"]
                if not all(field in evaluation for field in required_fields):
                    raise ValueError("Response missing required fields")
                
                # Ensure completeness_score is within bounds and unique
                score = float(evaluation["completeness_score"])
                
                # Check if this score has been used before for this document
                used_scores = [eval.get("completeness_score", 0) for eval in self.section_evaluations.values()]
                while score in used_scores:
                    # Adjust the score slightly to make it unique
                    score = max(0, min(100, score + (1 if score < 50 else -1)))
                
                evaluation["completeness_score"] = score
                
                # Ensure lists are actually lists
                if not isinstance(evaluation["missing_elements"], list):
                    evaluation["missing_elements"] = [str(evaluation["missing_elements"])]
                if not isinstance(evaluation["suggestions"], list):
                    evaluation["suggestions"] = [str(evaluation["suggestions"])]
                
                self.section_evaluations[section_name] = evaluation
                return evaluation
                
            except json.JSONDecodeError as je:
                print(f"Error parsing JSON response for section {section_name}: {je}")
                print("Raw response:", response_text)
                return {
                    "completeness_score": 50,
                    "missing_elements": ["Error parsing evaluation response"],
                    "suggestions": ["Please review this section manually", "Consider regenerating the evaluation"]
                }
                
        except Exception as e:
            print(f"Error evaluating section {section_name}: {e}")
            return {
                "completeness_score": 50,
                "missing_elements": ["Error during evaluation"],
                "suggestions": ["Please review this section manually"]
            }

    def save_documents(self,
                      documents: Dict[str, str],
                      output_directory: str) -> None:
        """
        Save generated documents as PDF, text, and markdown files.
        """
        output_path = Path(output_directory)
        output_path.mkdir(parents=True, exist_ok=True)

        for filename, content in documents.items():
            # Save as text file for backup
            txt_output_filename = f"{filename}_with_references.txt"
            txt_file_path = output_path / txt_output_filename
            with txt_file_path.open('w', encoding='utf-8') as f:
                f.write(content)

            # Save as markdown file
            md_output_filename = f"{filename}_with_references.md"
            md_file_path = output_path / md_output_filename
            with md_file_path.open('w', encoding='utf-8') as f:
                f.write(content)

            # Save as PDF
            pdf_output_filename = f"{filename}_with_references.pdf"
            pdf_file_path = output_path / pdf_output_filename
            self._convert_to_pdf(content, str(pdf_file_path))

            print(f"Saved document as PDF: {pdf_file_path}")
            print(f"Backup text file saved: {txt_file_path}")
            print(f"Markdown file saved: {md_file_path}")

    def _convert_to_pdf(self, content: str, output_path: str) -> None:
        """
        Convert text content to PDF using reportlab with reliable Unicode support.
        """
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            # Create document with generous margins
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Define styles
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(
                name='CustomBody',
                parent=styles['Normal'],
                fontSize=11,
                leading=14,
                spaceAfter=10
            ))
            styles.add(ParagraphStyle(
                name='CustomHeading',
                parent=styles['Heading1'],
                fontSize=14,
                leading=16,
                spaceAfter=12
            ))
            
            # Process content
            story = []
            lines = content.split('\n')
            
            for line in lines:
                if not line.strip():
                    story.append(Spacer(1, 12))
                    continue
                    
                # Handle special characters
                line = line.replace('"', '"').replace('"', '"')
                line = line.replace(''', "'").replace(''', "'")
                
                if line.startswith('#') or line.startswith('##'):
                    # Handle headers
                    story.append(Paragraph(line.replace('#', '').strip(), styles['CustomHeading']))
                else:
                    story.append(Paragraph(line, styles['CustomBody']))
            
            # Build PDF
            doc.build(story)
            print("PDF generated successfully with Unicode support")
            
        except Exception as e:
            print(f"Error creating PDF: {e}")
            print("Attempting even simpler PDF creation...")
            try:
                # Ultra-simple fallback using just basic text
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4
                
                c = canvas.Canvas(output_path, pagesize=A4)
                y = 800  # Start near top of page
                
                for line in content.split('\n'):
                    if not line.strip():
                        y -= 12
                        continue
                        
                    # Replace problematic characters
                    line = line.replace('"', '"').replace('"', '"')
                    line = line.replace(''', "'").replace(''', "'")
                    
                    if y < 50:  # Start new page if near bottom
                        c.showPage()
                        y = 800
                    
                    try:
                        c.drawString(50, y, line)
                    except:
                        c.drawString(50, y, line.encode('ascii', 'replace').decode())
                    y -= 12
                
                c.save()
                print("Created basic PDF as fallback")
            except Exception as e2:
                print(f"Fallback PDF creation failed: {e2}")
                print("Only text file will be available")

def load_document(file_path: str) -> str:
    """
    Load a document (text or markdown) from file.
    """
    path = Path(file_path)
    if path.exists():
        return path.read_text(encoding='utf-8')
    else:
        raise FileNotFoundError(f"{file_path} does not exist.")

def ai_assisted_edit(processor: MultiMediaProcessor, document_content: str, additional_context: str) -> str:
    """
    Use the AI model to refine the document content with additional context.
    """
    prompt = (
        "Below is the current document content:\n\n"
        f"{document_content}\n\n"
        "Additional context and instructions:\n"
        f"{additional_context}\n\n"
        "Please refine the document with the new context in mind. "
        "Provide an updated version that improves clarity and incorporates the additional details."
    )
    
    print(document_content)
    try:
        response = processor.client.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=[
                {"role": "system", "content": "You are an expert requirements document editor."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.7,
            top_p=0.9
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error during AI-assisted edit: {e}")
        return document_content  # Fallback to original content if an error occurs

def is_url(path_or_url: str) -> bool:
    """
    Check if the input is a URL.
    """
    return path_or_url.startswith("http://") or path_or_url.startswith("https://")

def is_blob_url(url: str) -> bool:
    """
    Check if the URL is an Azure Blob Storage URL.
    """
    return 'blob.core.windows.net' in url

def download_from_url(url: str) -> str:
    """
    Download a file from a URL and return the local path.
    For blob storage URLs, use the Azure SDK.
    """
    try:
        # Create a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_path = temp_file.name
        temp_file.close()
        
        if is_blob_url(url):
            # Parse the blob URL to get container and blob name
            parsed_url = urlparse(url)
            path_parts = parsed_url.path.split('/')
            container_name = path_parts[1]
            blob_name = '/'.join(path_parts[2:])
            
            # Get the connection string from environment or use default
            connection_string = os.environ.get('AZURE_STORAGE_CONNECTION_STRING', 
                                              "DefaultEndpointsProtocol=https;AccountName=barclaysstorage;AccountKey=your_account_key;EndpointSuffix=core.windows.net")
            
            # Create the BlobServiceClient
            blob_service_client = BlobServiceClient.from_connection_string(connection_string)
            
            # Get a reference to the container and blob
            container_client = blob_service_client.get_container_client(container_name)
            blob_client = container_client.get_blob_client(blob_name)
            
            # Download the blob
            with open(temp_path, "wb") as file:
                download_stream = blob_client.download_blob()
                file.write(download_stream.readall())
        else:
            # For regular URLs, use requests
            response = requests.get(url, stream=True)
            response.raise_for_status()
            
            with open(temp_path, "wb") as file:
                for chunk in response.iter_content(chunk_size=8192):
                    file.write(chunk)
        
        return temp_path
    except Exception as e:
        print(f"Error downloading from URL: {e}")
        return None

def extract_from_url(url: str, processor: MultiMediaProcessor) -> str:
    """
    Extract text from a document at a URL using Azure Document Intelligence.
    """
    try:
        # For blob storage URLs, use Form Recognizer directly on the URL
        if is_blob_url(url):
            print("Using Form Recognizer directly on blob URL...")
            poller = processor.form_recognizer_client.begin_analyze_document_from_url(
                "prebuilt-document", url
            )
            result = poller.result()
            
            # Extract all text content
            text_content = []
            
            # Get all paragraphs
            for paragraph in result.paragraphs:
                text_content.append(paragraph.content)
            
            # Get text from tables
            for table in result.tables:
                table_text = []
                for cell in table.cells:
                    table_text.append(cell.content)
                text_content.append(" | ".join(table_text))
            
            # Get key-value pairs
            for kv_pair in result.key_value_pairs:
                if kv_pair.key and kv_pair.value:
                    text_content.append(f"{kv_pair.key.content}: {kv_pair.value.content}")
            
            return "\n".join(text_content)
        
        # For regular URLs, download first then process
        else:
            # Download the file
            local_path = download_from_url(url)
            if not local_path:
                return ""
            
            # Use Form Recognizer directly for better extraction
            with open(local_path, "rb") as document:
                poller = processor.form_recognizer_client.begin_analyze_document(
                    "prebuilt-document", document
                )
                result = poller.result()
            
            # Extract all text content
            text_content = []
            
            # Get all paragraphs
            for paragraph in result.paragraphs:
                text_content.append(paragraph.content)
            
            # Get text from tables
            for table in result.tables:
                table_text = []
                for cell in table.cells:
                    table_text.append(cell.content)
                text_content.append(" | ".join(table_text))
            
            # Get key-value pairs
            for kv_pair in result.key_value_pairs:
                if kv_pair.key and kv_pair.value:
                    text_content.append(f"{kv_pair.key.content}: {kv_pair.value.content}")
            
            extracted_text = "\n".join(text_content)
            
            # Clean up the temporary file
            try:
                os.unlink(local_path)
            except:
                pass
                
            return extracted_text
    except Exception as e:
        print(f"Error extracting from URL with Form Recognizer: {e}")
        # Fallback to the processor's extract_text method
        try:
            if not is_blob_url(url):
                local_path = download_from_url(url)
                if local_path:
                    extracted_text = processor.extract_text(local_path)
                    # Clean up the temporary file
                    try:
                        os.unlink(local_path)
                    except:
                        pass
                    return extracted_text
        except Exception as e2:
            print(f"Fallback extraction also failed: {e2}")
        return ""

def upload_to_blob_storage(file_path: str, doc_type: str, doc_id: str) -> str:
    """
    Upload a PDF file to Azure Blob Storage and return a viewable URL.
    
    Args:
        file_path: Path to the PDF file
        doc_type: Type of document (BRD/SRS)
        doc_id: Document ID
        
    Returns:
        Viewable Blob URL
    """
    try:
        # Hard-coded connection string and container name
        connection_string = "DefaultEndpointsProtocol=https;AccountName=barclaysstorage;AccountKey=w7fJMnkuaZR4RX9LJbTRld8v90S6KDupj1BDHNZ1+Ch9Do7Et56nQKAd2HVXJqgZYEVbcGY/CGRj+AStE2NEXQ==;EndpointSuffix=core.windows.net"
        
        # Create the BlobServiceClient
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        
        # Use a fixed container name that's guaranteed to be valid
        container_name = "data"
        
        # Create the container if it doesn't exist
        try:
            container_client = blob_service_client.get_container_client(container_name)
            container_client.get_container_properties()
        except Exception:
            container_client = blob_service_client.create_container(container_name)
            print(f"Created container: {container_name}")
        
        # Create a valid blob name by replacing any invalid characters
        # Azure Blob Storage allows: alphanumeric characters, hyphens, and underscores
        safe_doc_type = re.sub(r'[^a-zA-Z0-9-]', '-', doc_type.lower())
        safe_doc_id = re.sub(r'[^a-zA-Z0-9-]', '-', doc_id)
        blob_name = f"{safe_doc_type}_{safe_doc_id}.pdf"
        
        # Upload the file with content type set to application/pdf
        with open(file_path, "rb") as data:
            blob_client = container_client.get_blob_client(blob_name)
            # Set content_type directly during upload
            blob_client.upload_blob(data, overwrite=True, content_type="application/pdf")
        
        # Get the blob URL
        blob_url = blob_client.url
        
        print(f"PDF uploaded to Blob Storage: {blob_url}")
        return blob_url
    except Exception as e:
        print(f"Error uploading to Blob Storage: {e}")
        return None

def store_document_in_cosmos(doc_type: str, content: str, project_id: str, pdf_path: str = None, requirement_doc_ids: List[str] = None) -> str:
    """
    Store the generated document in Cosmos DB.
    
    Args:
        doc_type: Type of document (BRD/SRS)
        content: Document content
        project_id: Project ID
        pdf_path: Path to the generated PDF file
        requirement_doc_ids: List of requirement document IDs
        
    Returns:
        Document ID
    """
    if not container:
        print("Cosmos DB container not available. Document will not be stored.")
        return None
        
    try:
        # Generate a unique document ID
        doc_id = f"std_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
        
        # Upload PDF to Blob Storage if available
        blob_url = None
        if pdf_path and os.path.exists(pdf_path):
            blob_url = upload_to_blob_storage(pdf_path, doc_type, doc_id)
        
        # Create document with the specified structure
        document = {
            "id": doc_id,
            "project_id": project_id,
            "doc_type": doc_type,
            "content": content,
            "version": 1,
            "status": "draft",
            "created_at": datetime.datetime.now().isoformat(),
            "requirement_doc_ids": requirement_doc_ids or [],
            "blob_url": blob_url
        }
        
        # Store in Cosmos DB using upsert_item to handle conflicts
        result = container.upsert_item(body=document)
        print(f"Document stored in Cosmos DB with ID: {doc_id}")
        
        if blob_url:
            print(f"PDF accessible at: {blob_url}")
            
        return doc_id
    except Exception as e:
        print(f"Error storing document in Cosmos DB: {e}")
        return None

def document_exists_in_cosmos(doc_id: str) -> bool:
    """
    Check if a document exists in Cosmos DB.
    
    Args:
        doc_id: Document ID to check
        
    Returns:
        True if document exists, False otherwise
    """
    if not container:
        print("Cosmos DB container not available. Cannot check if document exists.")
        return False
    
    try:
        # Try to read the document - if it succeeds, document exists
        container.read_item(item=doc_id, partition_key=doc_id)
        return True
    except exceptions.CosmosResourceNotFoundError:
        return False
    except Exception as e:
        print(f"Error checking if document exists: {e}")
        return False

def update_document_blob_url(doc_id: str, blob_url: str) -> bool:
    """
    Update the blob_url field of a document in Cosmos DB.
    
    Args:
        doc_id: Document ID
        blob_url: Blob URL to set
        
    Returns:
        True if successful, False otherwise
    """
    if not container:
        print("Cosmos DB container not available. Document will not be updated.")
        return False
        
    try:
        # Create a new document object with minimal required fields
        new_doc = {
            "id": doc_id,
            "blob_url": blob_url
        }
        
        # Use upsert to create or update the document
        if document_exists_in_cosmos(doc_id):
            # Get existing doc to preserve other fields
            existing_doc = container.read_item(item=doc_id, partition_key=doc_id)
            existing_doc["blob_url"] = blob_url
            existing_doc["version"] = existing_doc.get("version", 0) + 1
            container.upsert_item(body=existing_doc)
            print(f"Document updated in Cosmos DB with blob URL: {blob_url}")
            print(f"Document version: {existing_doc['version']}")
        else:
            # Create minimal document with blob URL
            new_doc.update({
                "project_id": "unknown",
                "doc_type": "unknown",
                "content": "Document content not available",
                "version": 1,
                "status": "draft",
                "created_at": datetime.datetime.now().isoformat(),
                "requirement_doc_ids": []
            })
            container.upsert_item(body=new_doc)
            print(f"New document created in Cosmos DB with ID: {doc_id} and blob URL: {blob_url}")
        
        return True
    except Exception as e:
        print(f"Error updating document in Cosmos DB: {e}")
        return False

def main():
    document_type = input("Enter document type (BRD/SRS): ").strip().upper()
    while document_type not in ['BRD', 'SRS']:
        print("Invalid document type. Please enter BRD or SRS.")
        document_type = input("Enter document type (BRD/SRS): ").strip().upper()

    business_domain = input("Enter business domain (e.g., healthcare, finance, retail) or press Enter to skip: ").strip()
    if not business_domain:
        business_domain = None
        
    project_id = input("Enter project ID: ").strip()
    if not project_id:
        project_id = f"proj_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
        print(f"Using generated project ID: {project_id}")

    print("\nEnter your requirements (you can provide both text input and document paths)")
    print("For text input, just type your requirements")
    print("For documents, enter the full path to the file or a URL")
    print("Supported formats: PDF, DOCX, XLSX, TXT, Images (JPG, PNG, etc.)")
    print("Type 'DONE' on a new line when finished\n")

    processor = MultiMediaProcessor()
    combined_content = []
    source_files = []

    while True:
        user_input = input("> ").strip()
        if user_input.lower() == 'done':
            if not combined_content and not source_files:
                print("No input provided. Exiting.")
                return
            break

        # Auto-detect if input is a URL
        if is_url(user_input):
            print(f"\nDetected URL: {user_input}")
            if is_blob_url(user_input):
                print("Azure Blob Storage URL detected")
            extracted_text = extract_from_url(user_input, processor)
            if extracted_text:
                print("\nExtracted content:")
                print("=" * 80)
                print(extracted_text)
                print("=" * 80)
                combined_content.append(extracted_text)
                source_files.append(user_input)
                print("\nURL processed and content appended successfully.")
            else:
                print("Failed to process URL.")
        # Check if input is a file path
        elif os.path.exists(user_input):
            print(f"\nProcessing document: {user_input}")
            extracted_text = processor.extract_text(user_input)
            if extracted_text:
                print("\nExtracted content:")
                print("=" * 80)
                print(extracted_text)
                print("=" * 80)
                combined_content.append(extracted_text)
                source_files.append(user_input)
                print("\nDocument processed and content appended successfully.")
            else:
                print("Failed to process document.")
        else:
            # Treat as text input
            combined_content.append(user_input)
            print("Text input recorded")

    output_directory = "./generated_documents_with_references"
    Path(output_directory).mkdir(parents=True, exist_ok=True)

    try:
        # Check if fpdf is installed
        try:
            import fpdf
        except ImportError:
            print("Installing required PDF library (fpdf)...")
            import subprocess
            subprocess.check_call(["pip", "install", "fpdf"])
            print("FPDF installed successfully!")

        # Create the combined content string
        all_content = "\n\n".join(combined_content)
        
        print("\nAll collected content:")
        print("=" * 80)
        print(all_content)
        print("=" * 80)

        proceed = input("\nProceed with document generation? (y/n): ").strip().lower()
        if proceed != 'y':
            print("Document generation cancelled.")
            return

        # Start spinner thread while generating document
        stop_spinner = threading.Event()
        spinner_thread = threading.Thread(target=spinner, args=(stop_spinner,))
        spinner_thread.start()
        
        # Store the content in processor for generation
        processor.original_source_content = all_content
        processor.original_source_names = [Path(f).name for f in source_files]
        processor.document_name = f"{document_type}_{datetime.datetime.now().strftime('%Y%m%d')}"
        processor.business_domain = business_domain

        # Generate document content in sections
        generated_doc = processor._generate_document_in_sections(
            document_type,
            all_content,
            processor.document_name,
            processor.original_source_names,
            business_domain
        )
        
        generated_documents = {processor.document_name: generated_doc}
        
        # Stop spinner
        stop_spinner.set()
        spinner_thread.join()

        # Get the document filename
        doc_filename = next(iter(generated_documents.keys()))
        
        # Save initial version of the document
        md_file = Path(output_directory) / f"{doc_filename}_with_references.md"
        with open(md_file, "w", encoding="utf-8") as f:
            f.write(generated_documents[doc_filename])

        # Print the generated content
        print("\nGenerated Document Content:")
        print("=" * 80)
        print(generated_documents[doc_filename])
        print("=" * 80)

        # Store the document in Cosmos DB without PDF
        doc_id = store_document_in_cosmos(
            doc_type=document_type,
            content=generated_documents[doc_filename],
            project_id=project_id
        )
        
        if doc_id:
            print(f"Document stored in Cosmos DB with ID: {doc_id}")
        else:
            print("Warning: Document was not stored in Cosmos DB. PDF generation will not update the document.")

        # Evaluate each section for completeness
        print("\nAnalyzing document sections for completeness...")
        print("\nSection Analysis:")
        print("=" * 80)
        
        for section, content in processor.generated_sections.items():
            evaluation = processor.evaluate_section_completeness(section, content)
            print(f"\nSection: {section}")
            print(f"Completeness Score: {evaluation['completeness_score']}%")
            print("Missing elements:")
            for item in evaluation["missing_elements"]:
                print(f"  - {item}")
            print("Suggestions:")
            for suggestion in evaluation["suggestions"]:
                print(f"  - {suggestion}")
            print("-" * 40)

        # Define the PDF file path
        pdf_file = Path(output_directory) / f"{doc_filename}_with_references.pdf"

        while True:
            print("\nOptions:")
            print("1. Edit document manually")
            print("2. Edit document with AI assistance")
            print("3. Generate final PDF")
            print("4. Exit")
            
            choice = input("\nEnter your choice (1-4): ").strip()
            
            if choice == "1":
                print(f"\nCurrent document content:")
                print("=" * 80)
                current_content = load_document(str(md_file))
                print(current_content)
                print("=" * 80)
                print(f"\nPlease edit the file: {md_file}")
                print("Press Enter when done editing...")
                input()
                
            elif choice == "2":
                additional_context = input("Enter additional context/instructions for the AI: ").strip()
                current_content = load_document(str(md_file))
                revised_content = ai_assisted_edit(processor, current_content, additional_context)
                
                print("\nAI suggested the following revision:\n")
                print("=" * 80)
                print(revised_content)
                print("=" * 80)
                
                accept = input("\nAccept changes? (y/n): ").strip().lower()
                if accept == 'y':
                    with open(md_file, 'w', encoding="utf-8") as f:
                        f.write(revised_content)
                    print("Document updated with AI revisions.")
                    
                    # Update the document in Cosmos DB if it exists
                    if doc_id and container:
                        try:
                            # Get the current document
                            doc = container.read_item(item=doc_id, partition_key=doc_id)
                            
                            # Update the content
                            doc["content"] = revised_content
                            doc["version"] += 1
                            
                            # Replace the document
                            container.replace_item(item=doc_id, body=doc)
                            print(f"Document updated in Cosmos DB with version {doc['version']}")
                        except Exception as e:
                            print(f"Error updating document in Cosmos DB: {e}")
                else:
                    print("No changes were made.")
                
            elif choice == "3":
                current_content = load_document(str(md_file))
                
                # Generate the final PDF
                processor._convert_to_pdf(current_content, str(pdf_file))
                print(f"\nFinal PDF generated: {pdf_file}")
                
                # Upload the final PDF to Blob Storage and update Cosmos DB
                if doc_id and container:
                    try:
                        # Upload the final PDF to Blob Storage
                        blob_url = upload_to_blob_storage(str(pdf_file), document_type, doc_id)
                        
                        if blob_url:
                            try:
                                # Get the current document
                                doc = container.read_item(item=doc_id, partition_key=doc_id)
                                
                                # Update the blob URL
                                doc["blob_url"] = blob_url
                                doc["version"] += 1
                                
                                # Replace the document
                                container.replace_item(item=doc_id, body=doc)
                                print(f"Final PDF uploaded to Blob Storage: {blob_url}")
                                print(f"Document updated in Cosmos DB with version {doc['version']}")
                            except exceptions.CosmosResourceNotFoundError:
                                print(f"Warning: Document with ID {doc_id} not found in Cosmos DB. Attempting to update blob URL.")
                                try:
                                    # Try to update the document using upsert
                                    new_doc = {
                                        "id": doc_id,
                                        "project_id": project_id,
                                        "doc_type": document_type,
                                        "content": current_content,
                                        "version": 1,
                                        "status": "draft",
                                        "created_at": datetime.datetime.now().isoformat(),
                                        "requirement_doc_ids": [],
                                        "blob_url": blob_url
                                    }
                                    # Use upsert_item instead of create_item to handle conflicts
                                    container.upsert_item(body=new_doc)
                                    print(f"Document updated in Cosmos DB with blob URL: {blob_url}")
                                except Exception as e:
                                    print(f"Error updating document in Cosmos DB: {e}")
                                    print(f"PDF was uploaded successfully and is accessible at: {blob_url}")
                    except Exception as e:
                        print(f"Error uploading final PDF to Blob Storage: {e}")
                else:
                    # If no doc_id or container, just upload the PDF
                    try:
                        # Generate a temporary doc_id for the PDF
                        temp_doc_id = f"std_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
                        blob_url = upload_to_blob_storage_from_memory(pdf_file.read(), document_type, temp_doc_id)
                        if blob_url:
                            print(f"Final PDF uploaded to Blob Storage: {blob_url}")
                            
                            # Try to create a new document in Cosmos DB
                            if container:
                                try:
                                    new_doc = {
                                        "id": temp_doc_id,
                                        "project_id": project_id,
                                        "doc_type": document_type,
                                        "content": current_content,
                                        "version": 1,
                                        "status": "draft",
                                        "created_at": datetime.datetime.now().isoformat(),
                                        "requirement_doc_ids": [],
                                        "blob_url": blob_url
                                    }
                                    # Use upsert_item instead of create_item to handle conflicts
                                    container.upsert_item(body=new_doc)
                                    print(f"New document created in Cosmos DB with ID: {temp_doc_id}")
                                except Exception as e:
                                    print(f"Error creating document in Cosmos DB: {e}")
                            else:
                                print("Note: Document was not stored in Cosmos DB as the container is not available.")
                    except Exception as e:
                        print(f"Error uploading final PDF to Blob Storage: {e}")
                
                break
                
            elif choice == "4":
                print("Exiting without generating final PDF.")
                break
                
            else:
                print("Invalid choice. Please select 1, 2, 3, or 4.")

    except Exception as e:
        print(f"An error occurred: {e}")
        import traceback
        print(traceback.format_exc())
        
    finally:
        # Clean up temporary text input file if it exists
        if source_files:
            try:
                for file in source_files:
                    if "text_input.txt" in file:
                        os.remove(file)
            except:
                pass

if __name__ == "__main__":
    # Check if we need to update a document's blob URL
    if len(sys.argv) > 2 and sys.argv[1] == "--update-blob-url":
        doc_id = sys.argv[2]
        blob_url = sys.argv[3]
        update_document_blob_url(doc_id, blob_url)
    else:
        main()

# Update the route to use the blueprint
@srs_brd_bp.route('/generate-standard-doc', methods=['POST'])
def generate_standard_doc():
    try:
        # Get request data
        data = request.get_json()
        project_id = data.get('project_id')
        doc_type = data.get('doc_type')  # 'SRS' or 'BRD'
        
        if not project_id or not doc_type:
            return jsonify({
                'success': False,
                'message': 'Missing required parameters: project_id and doc_type are required'
            }), 400

        # Get RequirementDocs container
        req_container = database.get_container_client("RequirementDocs")
        
        # Query to get all documents for this project
        query = f"SELECT * FROM RequirementDocs r WHERE r.project_id = '{project_id}'"
        items = list(req_container.query_items(
            query=query,
            enable_cross_partition_query=True
        ))

        if not items:
            return jsonify({
                'success': False,
                'message': 'No requirement documents found for this project'
            }), 404

        # Combine content from all documents
        combined_content = "\n\n".join([item.get('content', '') for item in items])
        
        # Initialize processor
        processor = MultiMediaProcessor()
        
        # Generate document based on type
        if doc_type.upper() == 'SRS':
            generated_doc = processor._generate_document_in_sections(
                'SRS',
                combined_content,
                f"SRS_{project_id}",
                [item.get('name', '') for item in items],
                None
            )
        elif doc_type.upper() == 'BRD':
            generated_doc = processor._generate_document_in_sections(
                'BRD',
                combined_content,
                f"BRD_{project_id}",
                [item.get('name', '') for item in items],
                None
            )
        else:
            return jsonify({
                'success': False,
                'message': 'Invalid document type. Must be either SRS or BRD'
            }), 400

        # Get StandardDocs container
        std_container = database.get_container_client("StandardDocs")
        
        # Check if document already exists
        query = f"SELECT * FROM StandardDocs s WHERE s.project_id = '{project_id}' AND s.template_type = '{doc_type.lower()}'"
        existing_docs = list(std_container.query_items(
            query=query,
            enable_cross_partition_query=True
        ))

        # Create document ID
        doc_id = str(uuid.uuid4())
        
        # Create temporary PDF file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            pdf_path = temp_pdf.name
            processor._convert_to_pdf(generated_doc, pdf_path)
            
            # Upload to blob storage
            blob_url = upload_to_blob_storage(pdf_path, doc_type.lower(), doc_id)
            
            # Clean up temporary file
            os.unlink(pdf_path)

        # Prepare document data
        doc_data = {
            'id': doc_id,
            'project_id': project_id,
            'template_type': doc_type.lower(),
            'context': generated_doc,
            'blob_url': blob_url,
            'pdf_filename': f"{doc_type}_{project_id}.pdf",
            'timestamp': datetime.utcnow().isoformat(),
            'version': 1
        }

        # Update or create document
        if existing_docs:
            # Update existing document
            existing_doc = existing_docs[0]
            doc_data['version'] = existing_doc.get('version', 1) + 1
            std_container.upsert_item(doc_data)
        else:
            # Create new document
            std_container.create_item(doc_data)

        return jsonify({
            'success': True,
            'message': f'{doc_type} generated successfully',
            'document': {
                'id': doc_id,
                'url': blob_url,
                'version': doc_data['version']
            }
        })

    except Exception as e:
        print(f"Error generating standard document: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Error generating document: {str(e)}'
        }), 500

@srs_brd_bp.route('/documents/<project_id>', methods=['GET'])
def get_documents(project_id):
    try:
        # Get StandardDocs container
        std_container = database.get_container_client("StandardDocs")
        
        # Query to get all documents for this project
        query = f"SELECT * FROM StandardDocs s WHERE s.project_id = '{project_id}'"
        items = list(std_container.query_items(
            query=query,
            enable_cross_partition_query=True
        ))

        return jsonify({
            'success': True,
            'documents': items
        })

    except Exception as e:
        print(f"Error fetching documents: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Error fetching documents: {str(e)}'
        }), 500

@srs_brd_bp.route('/generate-srs', methods=['POST'])
def generate_srs():
    try:
        # Get request data
        data = request.get_json()
        project_id = data.get('project_id')
        
        if not project_id:
            return jsonify({
                'success': False,
                'message': 'Missing required parameter: project_id'
            }), 400

        # Check if database is connected
        if not database:
            return jsonify({
                'success': False,
                'message': 'Database connection not available'
            }), 500

        # Get RequirementDocs container
        try:
            req_container = database.get_container_client("RequirementDocs")
            print(f"Successfully connected to RequirementDocs container for project {project_id}")
        except Exception as e:
            print(f"Error connecting to RequirementDocs container: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error connecting to RequirementDocs container: {str(e)}'
            }), 500
        
        # Query to get all documents for this project
        try:
            query = f"SELECT * FROM RequirementDocs r WHERE r.project_id = '{project_id}'"
            items = list(req_container.query_items(
                query=query,
                enable_cross_partition_query=True
            ))
            print(f"Found {len(items)} requirement documents for project {project_id}")
        except Exception as e:
            print(f"Error querying requirement documents: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error querying requirement documents: {str(e)}'
            }), 500

        if not items:
            return jsonify({
                'success': False,
                'message': 'No requirement documents found for this project'
            }), 404

        # Combine content from all documents
        try:
            # Extract content from each document, handling both string and dict content
            content_list = []
            for item in items:
                content = item.get('content', '')
                if isinstance(content, dict):
                    # If content is a dict, try to get the text content
                    content = content.get('text', '') or content.get('content', '') or str(content)
                elif not isinstance(content, str):
                    # If content is not a string, convert it to string
                    content = str(content)
                content_list.append(content)
            
            combined_content = "\n\n".join(content_list)
            print(f"Successfully combined content from {len(content_list)} documents")
        except Exception as e:
            print(f"Error combining document content: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error combining document content: {str(e)}'
            }), 500
        
        # Initialize processor
        try:
            processor = MultiMediaProcessor()
            print("Successfully initialized MultiMediaProcessor")
        except Exception as e:
            print(f"Error initializing MultiMediaProcessor: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error initializing document processor: {str(e)}'
            }), 500
        
        # Generate SRS document
        try:
            generated_doc = processor._generate_document_in_sections(
                'SRS',
                combined_content,
                f"SRS_{project_id}",
                [item.get('name', '') for item in items],
                None
            )
            print("Successfully generated SRS document content")
        except Exception as e:
            print(f"Error generating SRS document content: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error generating document content: {str(e)}'
            }), 500

        # Get StandardDocs container
        try:
            std_container = database.get_container_client("StandardDocs")
            print("Successfully connected to StandardDocs container")
        except Exception as e:
            print(f"Error connecting to StandardDocs container: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error connecting to StandardDocs container: {str(e)}'
            }), 500
        
        # Check if document already exists
        try:
            query = f"SELECT * FROM StandardDocs s WHERE s.project_id = '{project_id}' AND s.template_type = 'srs'"
            existing_docs = list(std_container.query_items(
                query=query,
                enable_cross_partition_query=True
            ))
            print(f"Found {len(existing_docs)} existing SRS documents")
        except Exception as e:
            print(f"Error checking for existing documents: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error checking for existing documents: {str(e)}'
            }), 500

        # Create document ID
        doc_id = str(uuid.uuid4())
        
        # Generate PDF in memory
        try:
            # Create a BytesIO object to store the PDF
            pdf_buffer = io.BytesIO()
            processor._convert_to_pdf(generated_doc, pdf_buffer)
            pdf_buffer.seek(0)
            print("Successfully generated PDF in memory")
            
            # Upload to blob storage directly from memory
            try:
                blob_url = upload_to_blob_storage_from_memory(pdf_buffer, 'srs', doc_id)
                print(f"Successfully uploaded PDF to blob storage: {blob_url}")
            except Exception as e:
                print(f"Error uploading to blob storage: {str(e)}")
                return jsonify({
                    'success': False,
                    'message': f'Error uploading document to storage: {str(e)}'
                }), 500
        except Exception as e:
            print(f"Error creating PDF: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error creating PDF: {str(e)}'
            }), 500

        # Prepare document data
        doc_data = {
            'id': doc_id,
            'project_id': project_id,
            'template_type': 'srs',
            'context': generated_doc,
            'blob_url': blob_url,
            'pdf_filename': f"SRS_{project_id}.pdf",
            'timestamp': datetime.utcnow().isoformat(),
            'version': 1
        }

        # Update or create document
        try:
            if existing_docs:
                # Update existing document
                existing_doc = existing_docs[0]
                doc_data['version'] = existing_doc.get('version', 1) + 1
                std_container.upsert_item(doc_data)
                print(f"Updated existing SRS document with version {doc_data['version']}")
            else:
                # Create new document
                std_container.create_item(doc_data)
                print("Created new SRS document")
        except Exception as e:
            print(f"Error storing document in database: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error storing document in database: {str(e)}'
            }), 500

        return jsonify({
            'success': True,
            'message': 'SRS generated successfully',
            'document': {
                'id': doc_id,
                'url': blob_url,
                'version': doc_data['version']
            }
        })

    except Exception as e:
        print(f"Error generating SRS: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Error generating SRS: {str(e)}'
        }), 500

@srs_brd_bp.route('/generate-brd', methods=['POST'])
def generate_brd():
    try:
        # Get request data
        data = request.get_json()
        project_id = data.get('project_id')
        
        if not project_id:
            return jsonify({
                'success': False,
                'message': 'Missing required parameter: project_id'
            }), 400

        # Check if database is connected
        if not database:
            return jsonify({
                'success': False,
                'message': 'Database connection not available'
            }), 500

        # Get RequirementDocs container
        try:
            req_container = database.get_container_client("RequirementDocs")
            print(f"Successfully connected to RequirementDocs container for project {project_id}")
        except Exception as e:
            print(f"Error connecting to RequirementDocs container: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error connecting to RequirementDocs container: {str(e)}'
            }), 500
        
        # Query to get all documents for this project
        try:
            query = f"SELECT * FROM RequirementDocs r WHERE r.project_id = '{project_id}'"
            items = list(req_container.query_items(
                query=query,
                enable_cross_partition_query=True
            ))
            print(f"Found {len(items)} requirement documents for project {project_id}")
        except Exception as e:
            print(f"Error querying requirement documents: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error querying requirement documents: {str(e)}'
            }), 500

        if not items:
            return jsonify({
                'success': False,
                'message': 'No requirement documents found for this project'
            }), 404

        # Combine content from all documents
        try:
            # Extract content from each document, handling both string and dict content
            content_list = []
            for item in items:
                content = item.get('content', '')
                if isinstance(content, dict):
                    # If content is a dict, try to get the text content
                    content = content.get('text', '') or content.get('content', '') or str(content)
                elif not isinstance(content, str):
                    # If content is not a string, convert it to string
                    content = str(content)
                content_list.append(content)
            
            combined_content = "\n\n".join(content_list)
            print(f"Successfully combined content from {len(content_list)} documents")
        except Exception as e:
            print(f"Error combining document content: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error combining document content: {str(e)}'
            }), 500
        
        # Initialize processor
        try:
            processor = MultiMediaProcessor()
            print("Successfully initialized MultiMediaProcessor")
        except Exception as e:
            print(f"Error initializing MultiMediaProcessor: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error initializing document processor: {str(e)}'
            }), 500
        
        # Generate BRD document
        try:
            generated_doc = processor._generate_document_in_sections(
                'BRD',
                combined_content,
                f"BRD_{project_id}",
                [item.get('name', '') for item in items],
                None
            )
            print("Successfully generated BRD document content")
        except Exception as e:
            print(f"Error generating BRD document content: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error generating document content: {str(e)}'
            }), 500

        # Get StandardDocs container
        try:
            std_container = database.get_container_client("StandardDocs")
            print("Successfully connected to StandardDocs container")
        except Exception as e:
            print(f"Error connecting to StandardDocs container: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error connecting to StandardDocs container: {str(e)}'
            }), 500
        
        # Check if document already exists
        try:
            query = f"SELECT * FROM StandardDocs s WHERE s.project_id = '{project_id}' AND s.template_type = 'brd'"
            existing_docs = list(std_container.query_items(
                query=query,
                enable_cross_partition_query=True
            ))
            print(f"Found {len(existing_docs)} existing BRD documents")
        except Exception as e:
            print(f"Error checking for existing documents: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error checking for existing documents: {str(e)}'
            }), 500

        # Create document ID
        doc_id = str(uuid.uuid4())
        
        # Generate PDF in memory
        try:
            # Create a BytesIO object to store the PDF
            pdf_buffer = io.BytesIO()
            processor._convert_to_pdf(generated_doc, pdf_buffer)
            pdf_buffer.seek(0)
            print("Successfully generated PDF in memory")
            
            # Upload to blob storage directly from memory
            try:
                blob_url = upload_to_blob_storage_from_memory(pdf_buffer, 'brd', doc_id)
                print(f"Successfully uploaded PDF to blob storage: {blob_url}")
            except Exception as e:
                print(f"Error uploading to blob storage: {str(e)}")
                return jsonify({
                    'success': False,
                    'message': f'Error uploading document to storage: {str(e)}'
                }), 500
        except Exception as e:
            print(f"Error creating PDF: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error creating PDF: {str(e)}'
            }), 500

        # Prepare document data
        doc_data = {
            'id': doc_id,
            'project_id': project_id,
            'template_type': 'brd',
            'context': generated_doc,
            'blob_url': blob_url,
            'pdf_filename': f"BRD_{project_id}.pdf",
            'timestamp': datetime.utcnow().isoformat(),
            'version': 1
        }

        # Update or create document
        try:
            if existing_docs:
                # Update existing document
                existing_doc = existing_docs[0]
                doc_data['version'] = existing_doc.get('version', 1) + 1
                std_container.upsert_item(doc_data)
                print(f"Updated existing BRD document with version {doc_data['version']}")
            else:
                # Create new document
                std_container.create_item(doc_data)
                print("Created new BRD document")
        except Exception as e:
            print(f"Error storing document in database: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error storing document in database: {str(e)}'
            }), 500

        return jsonify({
            'success': True,
            'message': 'BRD generated successfully',
            'document': {
                'id': doc_id,
                'url': blob_url,
                'version': doc_data['version']
            }
        })

    except Exception as e:
        print(f"Error generating BRD: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Error generating BRD: {str(e)}'
        }), 500

def upload_to_blob_storage_from_memory(content, doc_type: str, doc_id: str) -> str:
    """
    Upload a PDF from memory to Azure Blob Storage and return a viewable URL.
    
    Args:
        content: Bytes content of the PDF
        doc_type: Type of document (BRD/SRS)
        doc_id: Document ID
        
    Returns:
        Viewable Blob URL
    """
    try:
        # Hard-coded connection string and container name
        connection_string = "DefaultEndpointsProtocol=https;AccountName=barclaysstorage;AccountKey=w7fJMnkuaZR4RX9LJbTRld8v90S6KDupj1BDHNZ1+Ch9Do7Et56nQKAd2HVXJqgZYEVbcGY/CGRj+AStE2NEXQ==;EndpointSuffix=core.windows.net"
        
        # Create the BlobServiceClient
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        
        # Use a fixed container name that's guaranteed to be valid
        container_name = "data"
        
        # Create the container if it doesn't exist
        try:
            container_client = blob_service_client.get_container_client(container_name)
            container_client.get_container_properties()
        except Exception:
            container_client = blob_service_client.create_container(container_name)
            print(f"Created container: {container_name}")
        
        # Create a valid blob name by replacing any invalid characters
        safe_doc_type = re.sub(r'[^a-zA-Z0-9-]', '-', doc_type.lower())
        safe_doc_id = re.sub(r'[^a-zA-Z0-9-]', '-', doc_id)
        blob_name = f"{safe_doc_type}_{safe_doc_id}.pdf"
        
        # Upload the file with content type set to application/pdf
        blob_client = container_client.get_blob_client(blob_name)
        blob_client.upload_blob(content, overwrite=True, content_type="application/pdf")
        
        # Get the blob URL
        blob_url = blob_client.url
        
        print(f"PDF uploaded to Blob Storage: {blob_url}")
        return blob_url
    except Exception as e:
        print(f"Error uploading to Blob Storage: {e}")
        return None

@srs_brd_bp.route('/edit-document', methods=['POST'])
def edit_document():
    try:
        logger.info("=== Starting edit_document request ===")
        logger.info(f"Request headers: {dict(request.headers)}")
        
        data = request.get_json()
        logger.info(f"Request data: {data}")
        
        document_id = data.get('document_id')
        content = data.get('content')
        additional_context = data.get('additional_context', '')
        author = data.get('author', 'System User')  # Get author name from request
        change_summary = data.get('change_summary', 'Document updated')  # Get change summary

        if not document_id or not content:
            logger.error("Missing required parameters")
            return jsonify({
                'success': False,
                'error': 'Missing required parameters'
            }), 400

        # Get the existing document
        try:
            logger.info(f"Attempting to read document with ID: {document_id}")
            query = f"SELECT * FROM c WHERE c.id = '{document_id}'"
            items = list(cosmos_db.query_items(
                query=query,
                enable_cross_partition_query=True
            ))
            
            if not items:
                logger.error(f"No document found with ID: {document_id}")
                return jsonify({
                    'success': False,
                    'error': 'Document not found'
                }), 404
                
            document = items[0]
            partition_key = document.get('project_id')
            
            # Get the current version and increment it
            current_version = int(document.get('version', 1))
            new_version = current_version + 1
            logger.info(f"Current version: {current_version}, New version: {new_version}")

            # Create a unique ID for the new document version
            new_document_id = str(uuid.uuid4())
            
            # Check if there are any differences between old and new content
            if document['context'] == content and not additional_context and not change_summary:
                logger.info("No changes detected in document content")
                return jsonify({
                    'success': False,
                    'error': 'No changes detected in document content'
                }), 400

            # Create a new document entry with updated content and version
            new_document = {
                'id': new_document_id,
                'project_id': document['project_id'],
                'template_type': document['template_type'],
                'context': content,
                'version': new_version,
                'timestamp': document.get('timestamp'),  # Keep original creation timestamp
                'last_modified': datetime.now().isoformat(),
                'additional_context': additional_context,
                'parent_document_id': document_id,
                'author': author,  # Add author name
                'modified_by': author,  # Add who modified it
                'change_summary': change_summary,  # Add change summary
                'is_final': False,  # New versions are not final by default
                'is_latest_version': True  # Mark as the latest version
            }
            logger.info(f"New document data: {new_document}")

            # Generate new PDF for the new version
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                
                # Create a BytesIO object to store the PDF
                pdf_buffer = BytesIO()
                
                # Create document with generous margins
                doc = SimpleDocTemplate(
                    pdf_buffer,
                    pagesize=A4,
                    rightMargin=72,
                    leftMargin=72,
                    topMargin=72,
                    bottomMargin=72
                )
                
                # Define styles
                styles = getSampleStyleSheet()
                styles.add(ParagraphStyle(
                    name='CustomBody',
                    parent=styles['Normal'],
                    fontSize=11,
                    leading=14,
                    spaceAfter=10
                ))
                styles.add(ParagraphStyle(
                    name='CustomHeading',
                    parent=styles['Heading1'],
                    fontSize=14,
                    leading=16,
                    spaceAfter=12
                ))
                
                # Process content
                story = []
                lines = content.split('\n')
                
                # Add a version information header at the top of the document
                version_info = [
                    Paragraph(f"Document Version: {new_version}", styles['CustomHeading']),
                    Paragraph(f"Modified By: {author}", styles['CustomBody']),
                    Paragraph(f"Last Modified: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['CustomBody']),
                ]
                if change_summary:
                    version_info.append(Paragraph(f"Change Summary: {change_summary}", styles['CustomBody']))
                
                story.extend(version_info)
                story.append(Spacer(1, 20))  # Add space after the version info
                
                for line in lines:
                    if not line.strip():
                        story.append(Spacer(1, 12))
                        continue
                    
                    # Handle special characters and Unicode
                    line = line.replace('"', '"').replace('"', '"')
                    line = line.replace(''', "'").replace(''', "'")
                    
                    if line.startswith('#') or line.startswith('##'):
                        # Handle headers
                        story.append(Paragraph(line.replace('#', '').strip(), styles['CustomHeading']))
                    else:
                        story.append(Paragraph(line, styles['CustomBody']))
                
                # Build PDF
                doc.build(story)
                pdf_buffer.seek(0)
                logger.info("PDF generated successfully")
            except Exception as e:
                logger.error(f"Error generating PDF: {str(e)}")
                return jsonify({
                    'success': False,
                    'error': f'Error generating PDF: {str(e)}'
                }), 500

            # Upload new PDF to blob storage
            try:
                logger.info("Preparing to upload PDF to blob storage")
                # Get the PDF content as bytes
                pdf_content = pdf_buffer.getvalue()
                logger.info(f"PDF content size: {len(pdf_content)} bytes")
                
                # Upload to blob storage with version number in the filename
                blob_url = upload_to_blob_storage_from_memory(
                    pdf_content, 
                    f"{document['template_type']}_v{new_version}", 
                    new_document_id
                )
                if not blob_url:
                    raise Exception("Failed to upload PDF to blob storage")
                
                new_document['blob_url'] = blob_url
                logger.info(f"PDF uploaded successfully, blob URL: {blob_url}")
            except Exception as e:
                logger.error(f"Error uploading to blob storage: {str(e)}")
                return jsonify({
                    'success': False,
                    'error': f'Error uploading to blob storage: {str(e)}'
                }), 500

            # Store the new document in Cosmos DB
            try:
                logger.info("Storing new document in Cosmos DB")
                cosmos_db.create_item(body=new_document)
                logger.info("New document stored successfully")
            except Exception as e:
                logger.error(f"Error storing new document: {str(e)}")
                return jsonify({
                    'success': False,
                    'error': f'Error storing new document: {str(e)}'
                }), 500

            # Update the original document to mark it as not the latest version
            try:
                logger.info("Updating original document")
                document['is_latest_version'] = False
                cosmos_db.upsert_item(body=document)
                logger.info("Original document updated successfully")
            except Exception as e:
                logger.error(f"Error updating original document: {str(e)}")
                return jsonify({
                    'success': False,
                    'error': f'Error updating original document: {str(e)}'
                }), 500

            logger.info("=== Edit document request completed successfully ===")
            return jsonify({
                'success': True,
                'message': 'Document updated successfully',
                'document': {
                    'id': new_document['id'],
                    'version': new_version,
                    'blob_url': blob_url,
                    'timestamp': new_document['timestamp'],
                    'last_modified': new_document['last_modified'],
                    'modified_by': author,
                    'change_summary': change_summary
                }
            })

        except Exception as e:
            logger.error(f"=== Error in edit_document: {str(e)} ===")
            logger.error(f"Error type: {type(e)}")
            logger.error(f"Error traceback: {traceback.format_exc()}")
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    except Exception as e:
        logger.error(f"=== Error in edit_document: {str(e)} ===")
        logger.error(f"Error type: {type(e)}")
        logger.error(f"Error traceback: {traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@srs_brd_bp.route('/ai-assist', methods=['POST'])
def ai_assist():
    try:
        print("=== AI Assist Request Started ===")
        data = request.get_json()
        document_id = data.get('document_id')
        content = data.get('content')
        additional_context = data.get('additional_context', '')
        
        print(f"Document ID: {document_id}")
        print(f"Additional Context Length: {len(additional_context)}")
        
        if not document_id or not content:
            print("Error: Missing required parameters")
            return jsonify({
                'success': False,
                'message': 'Missing required parameters: document_id and content'
            }), 400

        # Initialize processor
        try:
            processor = MultiMediaProcessor()
            print("Successfully initialized MultiMediaProcessor")
        except Exception as e:
            print(f"Error initializing MultiMediaProcessor: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error initializing document processor: {str(e)}'
            }), 500

        # Get the document and its requirements
        try:
            std_container = database.get_container_client("StandardDocs")
            print(f"Looking for document with ID: {document_id}")
            
            # First query for the document
            query = f"SELECT * FROM c WHERE c.id = '{document_id}'"
            documents = list(std_container.query_items(
                query=query,
                enable_cross_partition_query=True
            ))
            
            if not documents:
                print(f"Document not found with ID: {document_id}")
                return jsonify({
                    'success': False,
                    'message': 'Document not found'
                }), 404
                
            doc = documents[0]
            doc_type = doc.get('template_type', 'srs').upper()
            project_id = doc.get('project_id')
            print(f"Found document: Type={doc_type}, Project ID={project_id}")

            # Get RequirementDocs container
            req_container = database.get_container_client("RequirementDocs")
            print(f"Successfully connected to RequirementDocs container for project {project_id}")
            
            # Query to get all documents for this project
            query = f"SELECT * FROM RequirementDocs r WHERE r.project_id = '{project_id}'"
            items = list(req_container.query_items(
                query=query,
                enable_cross_partition_query=True
            ))
            print(f"Found {len(items)} requirement documents for project {project_id}")

            if not items:
                return jsonify({
                    'success': False,
                    'message': 'No requirement documents found for this project'
                }), 404

            # Combine content from all documents using the combined_text field when available
            content_list = []
            for item in items:
                # Check if combined_text field exists and use it preferentially
                if 'combined_text' in item and item['combined_text']:
                    print(f"Using combined_text field for document {item.get('id', 'unknown')}")
                    content_list.append(item['combined_text'])
                else:
                    # Fall back to content field if combined_text is not available
                    req_content = item.get('content', '')
                    if isinstance(req_content, dict):
                        # If content is a dict, extract and join all values
                        dict_values = []
                        for key, value in req_content.items():
                            if isinstance(value, str):
                                dict_values.append(f"**{key}**: {value}")
                        if dict_values:
                            content_list.append("\n\n".join(dict_values))
                        else:
                            # Fallback to string representation if needed
                            content_list.append(str(req_content))
                    elif not isinstance(req_content, str):
                        # If content is not a string, convert it to string
                        content_list.append(str(req_content))
                    else:
                        # Content is already a string
                        content_list.append(req_content)
            
            combined_requirements = "\n\n".join(content_list)
            print(f"Successfully combined content from {len(content_list)} documents")
            print("=== COMBINED REQUIREMENTS START ===")
            print(combined_requirements[:1000] + "..." if len(combined_requirements) > 1000 else combined_requirements)
            print("=== COMBINED REQUIREMENTS END ===")

            print("=== CURRENT DOCUMENT CONTENT START ===")
            print(content[:1000] + "..." if len(content) > 1000 else content)
            print("=== CURRENT DOCUMENT CONTENT END ===")
            
            print("=== ADDITIONAL CONTEXT START ===")
            print(additional_context)
            print("=== ADDITIONAL CONTEXT END ===")

        except Exception as e:
            print(f"Error reading document or requirements: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return jsonify({
                'success': False,
                'message': f'Error reading document or requirements: {str(e)}'
            }), 404

        # Generate AI-assisted content
        try:
            # Create an enhanced, more focused prompt
            prompt = f"""You are an expert document reviewer tasked with revising a {doc_type} document according to specific instructions.

# TASK
Revise the provided document by STRICTLY following the additional instructions while ensuring alignment with the original requirements.

# ORIGINAL REQUIREMENTS
```
{combined_requirements}
```

# CURRENT DOCUMENT CONTENT
```
{content}
```

# ADDITIONAL INSTRUCTIONS - FOLLOW THESE PRECISELY
```
{additional_context}
```

# IMPORTANT GUIDELINES:
1. PRIORITIZE the additional instructions - they are critical client needs
2. Maintain the document's section structure and document type-specific formatting
3. Ensure all technical content remains accurate
4. Do not remove important details from the original document
5. Make sure each section addresses the corresponding requirements
6. If you add new content, ensure it logically fits within the document
7. Keep the same overall document organization and flow

Your revised document should result in a complete, well-structured {doc_type} that satisfies both the original requirements and the specific change requests in the additional instructions.
"""

            print("Sending enhanced prompt to AI model")
            
            # Use the client directly to get more control over the response
            response = processor.client.chat.completions.create(
                model=DEPLOYMENT_NAME,
                messages=[
                    {"role": "system", "content": f"You are an expert {doc_type} document editor with extensive experience in technical documentation."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,  # Lower temperature for more focus on instructions
                max_tokens=4000
            )
            
            improved_content = response.choices[0].message.content
            print(f"Successfully generated AI-assisted content ({len(improved_content)} characters)")
            
            return jsonify({
                'success': True,
                'message': 'AI assistance applied successfully',
                'content': improved_content
            })
        except Exception as e:
            print(f"Error generating AI-assisted content: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'Error generating AI-assisted content: {str(e)}'
            }), 500

    except Exception as e:
        print(f"Error in ai_assist: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Error getting AI assistance: {str(e)}'
        }), 500

@srs_brd_bp.route('/get-section-evaluations', methods=['POST'])
def get_section_evaluations():
    try:
        data = request.get_json()
        document_id = data.get('document_id')
        
        if not document_id:
            return jsonify({
                'success': False,
                'message': 'Missing required parameter: document_id'
            }), 400

        # Get the document
        std_container = database.get_container_client("StandardDocs")
        query = f"SELECT * FROM c WHERE c.id = '{document_id}'"
        documents = list(std_container.query_items(
            query=query,
            enable_cross_partition_query=True
        ))
        
        if not documents:
            return jsonify({
                'success': False,
                'message': 'Document not found'
            }), 404
            
        doc = documents[0]
        content = doc.get('context', '')
        doc_type = doc.get('template_type', 'srs').upper()

        # Initialize processor
        processor = MultiMediaProcessor()
        
        # Parse sections from content
        sections = {}
        current_section = None
        current_content = []
        
        for line in content.split('\n'):
            line = line.strip()
            if line.startswith('## '):
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = line[3:].strip()
                current_content = []
            elif current_section:
                current_content.append(line)
        
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)

        # Evaluate each section
        evaluations = {}
        for section_name, section_content in sections.items():
            evaluation = processor.evaluate_section_completeness(section_name, section_content)
            evaluations[section_name] = {
                'completeness_score': evaluation['completeness_score'],
                'missing_elements': evaluation['missing_elements']
            }

        return jsonify({
            'success': True,
            'evaluations': evaluations
        })

    except Exception as e:
        print(f"Error getting section evaluations: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Error getting section evaluations: {str(e)}'
        }), 500

@srs_brd_bp.route('/test', methods=['GET'])
def test_route():
    logger.info("Test route accessed")
    return jsonify({
        'success': True,
        'message': 'Test route is working'
    })

@srs_brd_bp.route('/set-document-final', methods=['POST'])
def set_document_final():
    try:
        data = request.get_json()
        document_id = data.get('document_id')
        author = data.get('author')
        is_final = data.get('is_final', True)  # Default to True if not specified

        if not document_id:
            return jsonify({'success': False, 'error': 'Document ID is required'}), 400

        # First get the document to get its partition key
        query = f"SELECT * FROM c WHERE c.id = '{document_id}'"
        items = list(cosmos_db.query_items(query=query, enable_cross_partition_query=True))
        
        if not items:
            return jsonify({'success': False, 'error': 'Document not found'}), 404
            
        document = items[0]
        partition_key = document['project_id']

        # Update the document
        document['is_final'] = is_final
        if is_final:
            document['finalized_by'] = author
            document['finalized_at'] = datetime.utcnow().isoformat()
        else:
            document['finalized_by'] = None
            document['finalized_at'] = None

        # Save the updated document
        cosmos_db.upsert_item(body=document)

        return jsonify({
            'success': True,
            'document': {
                'id': document['id'],
                'is_final': document['is_final'],
                'finalized_by': document['finalized_by'],
                'finalized_at': document['finalized_at']
            }
        })

    except Exception as e:
        logger.error(f"Error setting document final status: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@srs_brd_bp.route('/document-versions/<project_id>/<template_type>', methods=['GET'])
def get_document_versions(project_id, template_type):
    """
    Retrieve all versions of a specific document type for a project
    """
    try:
        logger.info(f"Fetching versions for {template_type} documents in project {project_id}")
        
        # Query for all documents of the specified template type in this project
        query = f"SELECT * FROM c WHERE c.project_id = '{project_id}' AND c.template_type = '{template_type}'"
        items = list(cosmos_db.query_items(
            query=query,
            enable_cross_partition_query=True
        ))
        
        if not items:
            logger.info(f"No {template_type} documents found for project {project_id}")
            return jsonify({
                'success': False,
                'error': f'No {template_type} documents found'
            }), 404
        
        # Sort documents by version
        items.sort(key=lambda x: x.get('version', 1), reverse=True)
        
        # Format the response
        versions = []
        for item in items:
            version_data = {
                'id': item['id'],
                'version': item.get('version', 1),
                'timestamp': item.get('timestamp'),
                'last_modified': item.get('last_modified', item.get('timestamp')),
                'modified_by': item.get('modified_by'),
                'is_final': item.get('is_final', False),
                'finalized_by': item.get('finalized_by'),
                'finalized_at': item.get('finalized_at'),
                'blob_url': item.get('blob_url'),
                'change_summary': item.get('change_summary'),
                'is_latest_version': item.get('is_latest_version', False)
            }
            versions.append(version_data)
        
        logger.info(f"Found {len(versions)} versions of {template_type} documents")
        return jsonify({
            'success': True,
            'versions': versions
        })
        
    except Exception as e:
        logger.error(f"Error retrieving document versions: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': f'Failed to retrieve document versions: {str(e)}'
        }), 500