import io
import base64
import json
from pathlib import Path
import re
from openai import AzureOpenAI
import PyPDF2
import pandas as pd
import docx
from PIL import Image
import requests
import os
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
import httpx
import tempfile
from azure.cosmos import CosmosClient, exceptions
from datetime import datetime
from flask import Blueprint, request, jsonify
import traceback

# Initialize Flask Blueprint
srs_brd_to_stories_bp = Blueprint('srs_brd_to_stories', __name__)

# Azure OpenAI Configuration
AZURE_ENDPOINT = "https://suraj-m9lgdbv9-eastus2.cognitiveservices.azure.com/"
AZURE_API_KEY = "75PVa3SAy9S2ZR590gZesyTNDMZtb3Oa5EdHlRbqWeQ89bmoOGl4JQQJ99BDACHYHv6XJ3w3AAAAACOGUP8d"
AZURE_API_VERSION = "2024-02-15-preview"
DEPLOYMENT_NAME = "gpt-4v"  # Make sure this is your vision model deployment name
TEXT_DEPLOYMENT_NAME = "gpt-4"  # Your text model deployment name

# Azure Document Intelligence Configuration
FORM_ENDPOINT = "https://barclaysform.cognitiveservices.azure.com/"
FORM_KEY = "63spGg0VYFV0kWZB3nmsFDp8yEbi40zmEnCvIl6D8Seih4YyLsp9JQQJ99BDACYeBjFXJ3w3AAALACOGh5hu"

# Azure Cosmos DB Configuration
COSMOS_ENDPOINT = "https://barclaysdb.documents.azure.com:443/"
COSMOS_KEY = "ercuc7wFNt4RPsxcx2QTzKP8AhKUDnjJrt0mpZrW2Yi2IQGAa7wDrEbhKRHsurGu0P7GuGny4IZRACDbtecfNQ=="
COSMOS_DATABASE = "RequirementsDB"  # Changed to match existing database
COSMOS_CONTAINER = "UserStories"    # Changed to match existing container

# Initialize clients
openai_client = AzureOpenAI(
    api_version=AZURE_API_VERSION,
    api_key=AZURE_API_KEY,
    azure_endpoint=AZURE_ENDPOINT
)

document_client = DocumentAnalysisClient(
    endpoint=FORM_ENDPOINT, 
    credential=AzureKeyCredential(FORM_KEY)
)

# Initialize Cosmos DB client with error handling
try:
    cosmos_client = CosmosClient(COSMOS_ENDPOINT, COSMOS_KEY)
    database = cosmos_client.get_database_client(COSMOS_DATABASE)
    container = database.get_container_client(COSMOS_CONTAINER)
    print("Successfully connected to Cosmos DB")
except Exception as e:
    print(f"Error connecting to Cosmos DB: {e}")
    cosmos_client = None
    database = None
    container = None

def is_url(path_or_url: str) -> bool:
    """
    Check if the input is a URL.
    """
    return path_or_url.startswith("http://") or path_or_url.startswith("https://")

def download_from_url(url: str) -> Path:
    """
    Download a file from a URL and return the local path.
    """
    try:
        response = httpx.get(url)
        response.raise_for_status()
        
        # Create a temporary file with the appropriate extension
        suffix = Path(url).suffix
        if not suffix:
            # Default to .pdf if no extension is found
            suffix = ".pdf"
            
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        temp_file.write(response.content)
        temp_file.close()
        
        return Path(temp_file.name)
    except Exception as e:
        print(f"Error downloading file from URL: {e}")
        raise

def extract_with_document_intelligence(path: Path) -> str:
    """
    Extract text from any document using Azure Document Intelligence.
    This provides better extraction for complex documents with tables, forms, etc.
    """
    try:
        with open(path, "rb") as document:
            poller = document_client.begin_analyze_document(
                "prebuilt-document", document
            )
            result = poller.result()

            # Extract all text content
            text_content = []
            
            # Get all paragraphs
            for paragraph in result.paragraphs:
                text_content.append(paragraph.content)
            
            # Get text from tables
            for table in result.tables:
                table_text = []
                for cell in table.cells:
                    table_text.append(cell.content)
                text_content.append(" | ".join(table_text))
            
            # Get key-value pairs
            for kv_pair in result.key_value_pairs:
                if kv_pair.key and kv_pair.value:
                    text_content.append(f"{kv_pair.key.content}: {kv_pair.value.content}")

            return "\n".join(text_content)
    except Exception as e:
        print(f"Error in Document Intelligence processing: {e}")
        return ""

def extract_from_url(url: str) -> str:
    """
    Extract text from a document URL using Azure Document Intelligence.
    """
    try:
        # Download the file from the URL
        local_path = download_from_url(url)
        
        # Extract text using Document Intelligence
        text = extract_with_document_intelligence(local_path)
        
        # Clean up the temporary file
        try:
            os.unlink(local_path)
        except:
            pass
            
        return text
    except Exception as e:
        print(f"Error extracting text from URL: {e}")
        return ""

def load_file_text(file_path: str) -> str:
    """
    Extract text from various file types or URLs.
    First attempts Azure Document Intelligence, then falls back to specific extractors.
    """
    # Check if the input is a URL
    if is_url(file_path):
        print(f"Processing document from URL: {file_path}")
        return extract_from_url(file_path)
    
    # Process local file
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"{file_path} does not exist.")
    
    suffix = path.suffix.lower()
    
    try:
        # First try Azure Document Intelligence for supported formats
        if suffix in ['.pdf', '.docx', '.doc', '.jpeg', '.jpg', '.png', '.bmp', '.tiff', '.tif']:
            print("Using Azure Document Intelligence for extraction...")
            text = extract_with_document_intelligence(path)
            if text.strip():
                return text
            print("Falling back to specific extractors...")
        
        # Fall back to specific extractors
        if suffix in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            return process_image(path)
        elif suffix == '.pdf':
            return extract_pdf_text(path)
        elif suffix == '.docx':
            return extract_docx_text(path)
        elif suffix in ['.xls', '.xlsx']:
            return extract_excel_text(path)
        elif suffix == '.txt':
            return path.read_text(encoding="utf-8")
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return ""

def process_image(path: Path) -> str:
    """
    Process an image file: resize if needed, encode to base64,
    and have the model describe it.
    """
    with Image.open(path) as img:
        max_size = (800, 800)
        img.thumbnail(max_size)
        buffer = io.BytesIO()
        fmt = img.format if img.format else "JPEG"
        img.save(buffer, format=fmt)
        buffer.seek(0)
    
    image_data = base64.b64encode(buffer.getvalue()).decode('utf-8')
    
    try:
        response = openai_client.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Thoroughly describe the contents of this image. Extract any visible text and provide a comprehensive analysis."},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_data}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=1000
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error processing image: {e}")
        return ""

def extract_pdf_text(path: Path) -> str:
    """Extract text from a PDF file."""
    with path.open("rb") as file:
        reader = PyPDF2.PdfReader(file)
        texts = [page.extract_text() for page in reader.pages if page.extract_text()]
        return "\n".join(texts)

def extract_docx_text(path: Path) -> str:
    """Extract text from a DOCX file."""
    doc = docx.Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

def extract_excel_text(path: Path) -> str:
    """Extract text from an Excel file."""
    df = pd.read_excel(path)
    return df.to_string()

def extract_user_stories_json(document: str) -> list:
    """
    Generate functional user stories from a technical document using Azure OpenAI.
    Returns a list of user story dictionaries.
    """
    prompt = (
        "Extract user stories from the following technical requirement document and output ONLY a valid JSON array. "
        "Each element in the array must be a JSON object with the keys: 'title', 'user_role', 'goal', 'benefit', "
        "'acceptance_criteria', 'priority', and 'document'. "
        "The 'priority' field should be one of: 'Must Have', 'Should Have', 'Could Have', or 'Won't Have'. "
        "Do not output anything else besides the JSON array.\n\n"
        "Document:\n\n{document}\n\n"
        "Guidelines:\n"
        "- Extract only functional user stories.\n"
        "- Ensure clarity and avoid technical jargon.\n"
        "- Group related stories if applicable.\n"
        "- Use the MoSCoW prioritization framework for the priority field."
    ).format(document=document)
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",  # Use your text model deployment
            messages=[
                {"role": "system", "content": "You are a requirements analyst that creates user stories from technical documents. Always respond with valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=3000
        )
        
        cleaned_text = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
        
        # Attempt to parse the JSON
        try:
            stories = json.loads(cleaned_text)
            return stories
        except json.JSONDecodeError:
            # Attempt to extract the JSON array substring
            match = re.search(r'\[.*\]', cleaned_text, re.DOTALL)
            if match:
                json_text = match.group(0)
                try:
                    stories = json.loads(json_text)
                    return stories
                except json.JSONDecodeError as e:
                    print("Warning: Failed to parse JSON after substring extraction:", e)
                    print("Extracted substring:")
                    print(json_text)
                    # As a final hack, if missing a closing bracket, add it.
                    if not json_text.strip().endswith("]"):
                        fixed_text = json_text.strip() + "]"
                        try:
                            stories = json.loads(fixed_text)
                            return stories
                        except Exception as e:
                            print("Warning: Failed to fix JSON by appending closing bracket:", e)
                    return []
            else:
                print("Warning: No JSON array found in the response.")
                print("Raw cleaned response:")
                print(cleaned_text)
                return []
    except Exception as e:
        print(f"Error generating user stories: {e}")
        return []

def push_story_to_jira(story: dict):
    """
    Push a user story to Jira using direct credentials.
    """
    # Direct JIRA credentials
    jira_domain = "prathamgadkari.atlassian.net"
    jira_email = "prathamgadkari@gmail.com"
    jira_api_token = "ATATT3xFfGF0LXSWvbAlLm1lVXyxSroefrBXONpTLDg5mzfknXUUxTYAjQ0u59wKKGh3ObCEnducPVGqRwCLoxDu8oc4xx3aBAHj6A9Tzuj2OqiUszHVo3UvYrpmblYC2xHttFxo5ULieeRE2LNIfR3w_l2YNJTvACQ_zmBnkt6Tjvenqyu2pEM=99E9EA51"
    project_key = "SCRUM"

    # Validate required fields
    required_fields = ['title', 'user_role', 'goal', 'benefit']
    missing_fields = [field for field in required_fields if not story.get(field)]
    if missing_fields:
        print(f"Error: Missing required fields for Jira: {missing_fields}")
        return None

    jira_url = f"https://{jira_domain}/rest/api/2/issue"
    auth = (jira_email, jira_api_token)
    headers = {"Content-Type": "application/json"}
    
    # Format the description with all story details
    description = f"""
*User Role:* {story.get('user_role', '')}
*Goal:* {story.get('goal', '')}
*Benefit:* {story.get('benefit', '')}
*Acceptance Criteria:* {story.get('acceptance_criteria', '')}
*Priority:* {story.get('priority', '')}
*Source Document:* {story.get('source_doc_type', '')}
    """.strip()
    
    # Map our priorities to Jira priorities
    priority_mapping = {
        'Must Have': 'Highest',
        'Should Have': 'High',
        'Could Have': 'Medium',
        'Won\'t Have': 'Low'
    }
    
    # Get the mapped priority or default to Medium
    jira_priority = priority_mapping.get(story.get('priority', ''), 'Medium')
    
    payload = {
        "fields": {
            "project": {"key": project_key},
            "summary": story.get('title', 'Untitled Story'),
            "description": description,
            "issuetype": {"name": "Story"}
        }
    }

    try:
        response = requests.post(jira_url, auth=auth, headers=headers, data=json.dumps(payload))
        if response.status_code == 201:
            print(f"Story '{story.get('title')}' successfully pushed to Jira.")
            return response.json().get('key')  # Return the Jira issue key
        else:
            print(f"Failed to push story '{story.get('title')}' to Jira. Status code: {response.status_code}. Response: {response.text}")
            return None
    except Exception as e:
        print(f"Error pushing story '{story.get('title')}' to Jira: {e}")
        return None

def store_story_in_cosmos(story: dict, project_id: str, standard_doc_id: str) -> str:
    """
    Store a user story in Cosmos DB following the specified structure.
    Returns the story ID.
    """
    if not container:
        print("Cosmos DB connection not available. Skipping story storage.")
        return None
        
    try:
        # Generate a unique story ID
        story_id = f"story_{project_id}_{standard_doc_id}_{int(datetime.now().timestamp())}"
        
        # Create the story document with all fields from the JSON structure
        story_doc = {
            "id": story_id,
            "project_id": project_id,
            "standard_doc_id": standard_doc_id,
            "title": story.get("title", "Untitled Story"),
            "user_role": story.get("user_role", ""),
            "goal": story.get("goal", ""),
            "benefit": story.get("benefit", ""),
            "acceptance_criteria": story.get("acceptance_criteria", ""),
            "priority": story.get("priority", "Must Have"),
            "document": story.get("document", ""),
            "jira_issue_id": story.get("jira_issue_id", ""),
            "status": "To Do",
            "created_at": datetime.utcnow().isoformat(),
        }
        
        # Create the item in Cosmos DB
        container.create_item(body=story_doc)
        print(f"Story '{story_doc['title']}' successfully stored in Cosmos DB.")
        return story_id
        
    except exceptions.CosmosHttpResponseError as e:
        print(f"Cosmos DB error storing story: {e}")
        return None
    except Exception as e:
        print(f"Error storing story in Cosmos DB: {e}")
        return None

@srs_brd_to_stories_bp.route('/generate-stories', methods=['POST'])
def generate_stories_from_docs():
    try:
        print("\n=== Starting generate-stories request ===")
        print("Request headers:", request.headers)
        print("Request data:", request.get_json())
        
        # Get request data
        data = request.get_json()
        project_id = data.get('project_id')
        doc_types = data.get('doc_types', [])  # ['SRS', 'BRD', 'BOTH']
        
        print(f"Project ID: {project_id}")
        print(f"Document types: {doc_types}")
        
        if not project_id:
            print("Error: Project ID is missing")
            return jsonify({'error': 'Project ID is required'}), 400
            
        if not doc_types:
            print("Error: Document types are missing")
            return jsonify({'error': 'Document types are required'}), 400

        # Get database and containers
        print("Connecting to Cosmos DB...")
        database = cosmos_client.get_database_client(COSMOS_DATABASE)
        standard_docs_container = database.get_container_client("StandardDocs")
        user_stories_container = database.get_container_client("UserStories")
        
        # Query for final documents
        query = f"SELECT * FROM c WHERE c.project_id = '{project_id}' AND c.is_final = true"
        if 'BOTH' not in doc_types:
            # Map SRS/BRD to template_type values
            doc_type_mapping = {
                'SRS': 'srs',
                'BRD': 'brd'
            }
            doc_type_filter = " OR ".join([f"c.template_type = '{doc_type_mapping[doc_type]}'" for doc_type in doc_types])
            query += f" AND ({doc_type_filter})"
        
        print(f"Querying documents with: {query}")
        items = list(standard_docs_container.query_items(
            query=query,
            enable_cross_partition_query=True
        ))
        
        print(f"Found {len(items)} documents")
        if not items:
            print("No documents found matching criteria")
            return jsonify({'error': 'No final documents found for the specified criteria'}), 404
            
        # Process each document and generate stories
        all_stories = []
        for doc in items:
            try:
                print(f"\nProcessing document: {doc.get('id')}")
                # Extract text from document context
                document_text = doc.get('context', '')
                if not document_text:
                    print("Document has no context, skipping")
                    continue
                
                # Print the document context used for generating stories
                print("\n=== Document Context Used for Story Generation ===")
                print(f"Document ID: {doc.get('id')}")
                print(f"Document Title: {doc.get('title', 'No Title')}")
                print(f"Document Type: {doc.get('template_type', 'Unknown')}")
                print("Context Content (first 500 chars):")
                print(document_text[:500] + "..." if len(document_text) > 500 else document_text)
                print("=== End of Document Context Sample ===\n")
                    
                # Generate user stories
                print("Generating user stories...")
                stories = extract_user_stories_json(document_text)
                print(f"Generated {len(stories)} stories")
                
                # Store stories in UserStories container
                for story in stories:
                    # Generate a unique story ID
                    story_id = f"story_{project_id}_{doc['id']}_{int(datetime.now().timestamp())}"
                    
                    # Create the story document
                    story_doc = {
                        "id": story_id,
                        "project_id": project_id,
                        "standard_doc_id": doc['id'],
                        "title": story.get("title", "Untitled Story"),
                        "user_role": story.get("user_role", ""),
                        "goal": story.get("goal", ""),
                        "benefit": story.get("benefit", ""),
                        "acceptance_criteria": story.get("acceptance_criteria", ""),
                        "priority": story.get("priority", "Must Have"),
                        "document": doc.get("title", ""),
                        "status": "To Do",
                        "created_at": datetime.utcnow().isoformat(),
                        "source_doc_type": doc.get("template_type", "Unknown"),
                        "jira_issue_id": ""
                    }
                    
                    print(f"Storing story: {story_doc['title']}")
                    try:
                        # Store in UserStories container
                        user_stories_container.create_item(body=story_doc)
                        print(f"Successfully stored story with ID: {story_id}")
                        story['id'] = story_id
                        all_stories.append(story)
                    except Exception as e:
                        print(f"Error storing story in database: {str(e)}")
                        continue
                        
            except Exception as e:
                print(f"Error processing document {doc['id']}: {str(e)}")
                continue
                
        print(f"\nSuccessfully generated {len(all_stories)} stories")
        return jsonify({
            'success': True,
            'count': len(all_stories),
            'stories': all_stories
        })
        
    except Exception as e:
        print(f"\nError in generate-stories route: {str(e)}")
        print("Stack trace:", traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@srs_brd_to_stories_bp.route('/stories', methods=['GET'])
def get_all_stories():
    try:
        print("\n=== Starting get_all_stories request ===")
        database = cosmos_client.get_database_client(COSMOS_DATABASE)
        user_stories_container = database.get_container_client("UserStories")
        
        # Query all stories
        query = "SELECT * FROM c"
        items = list(user_stories_container.query_items(
            query=query,
            enable_cross_partition_query=True
        ))
        
        print(f"Found {len(items)} stories")
        return jsonify({
            'success': True,
            'stories': items
        })
    except Exception as e:
        print(f"Error fetching stories: {str(e)}")
        print("Stack trace:", traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@srs_brd_to_stories_bp.route('/stories/<project_id>', methods=['GET'])
def get_project_stories(project_id):
    try:
        print("\n=== Starting get_project_stories request ===")
        print(f"Project ID: {project_id}")
        
        database = cosmos_client.get_database_client(COSMOS_DATABASE)
        user_stories_container = database.get_container_client("UserStories")
        
        # Query stories for the project
        query = f"SELECT * FROM c WHERE c.project_id = '{project_id}'"
        items = list(user_stories_container.query_items(
            query=query,
            enable_cross_partition_query=True
        ))
        
        print(f"Found {len(items)} stories for project {project_id}")
        return jsonify({
            'success': True,
            'stories': items
        })
    except Exception as e:
        print(f"Error fetching project stories: {str(e)}")
        print("Stack trace:", traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@srs_brd_to_stories_bp.route('/stories/<story_id>', methods=['GET'])
def get_story(story_id):
    try:
        print("\n=== Starting get_story request ===")
        print(f"Story ID: {story_id}")
        
        database = cosmos_client.get_database_client(COSMOS_DATABASE)
        user_stories_container = database.get_container_client("UserStories")
        
        # Get the story
        story = user_stories_container.read_item(item=story_id, partition_key=story_id)
        
        print("Story found:", story)
        return jsonify({
            'success': True,
            'story': story
        })
    except exceptions.CosmosResourceNotFoundError:
        print(f"Story not found with ID: {story_id}")
        return jsonify({'error': 'Story not found'}), 404
    except Exception as e:
        print(f"Error fetching story: {str(e)}")
        print("Stack trace:", traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@srs_brd_to_stories_bp.route('/stories/<story_id>', methods=['DELETE'])
def delete_story(story_id):
    try:
        print("\n=== Starting delete_story request ===")
        print(f"Story ID: {story_id}")
        
        database = cosmos_client.get_database_client(COSMOS_DATABASE)
        user_stories_container = database.get_container_client("UserStories")
        
        # First find the story using a query
        query = f"SELECT * FROM c WHERE c.id = '{story_id}'"
        items = list(user_stories_container.query_items(
            query=query,
            enable_cross_partition_query=True
        ))
        
        if not items:
            print(f"Story not found with ID: {story_id}")
            return jsonify({'error': 'Story not found'}), 404
            
        story = items[0]
        print("Story found:", story)
        
        try:
            # Delete the story using its ID and partition key
            user_stories_container.delete_item(
                item=story_id,
                partition_key=story.get('project_id', story_id)  # Use project_id as partition key
            )
            print("Story deleted successfully")
            return jsonify({
                'success': True,
                'message': 'Story deleted successfully'
            })
        except Exception as e:
            print(f"Error deleting story: {str(e)}")
            print("Stack trace:", traceback.format_exc())
            return jsonify({'error': str(e)}), 500
            
    except Exception as e:
        print(f"Error in delete_story route: {str(e)}")
        print("Stack trace:", traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@srs_brd_to_stories_bp.route('/stories/<story_id>/status', methods=['PUT'])
def update_story_status(story_id):
    try:
        print("\n=== Starting update_story_status request ===")
        print(f"Story ID: {story_id}")
        
        data = request.get_json()
        new_status = data.get('status')
        
        if not new_status:
            return jsonify({'error': 'Status is required'}), 400
            
        database = cosmos_client.get_database_client(COSMOS_DATABASE)
        user_stories_container = database.get_container_client("UserStories")
        
        # Get the story
        story = user_stories_container.read_item(item=story_id, partition_key=story_id)
        
        # Update the status
        story['status'] = new_status
        user_stories_container.replace_item(item=story_id, body=story)
        
        print(f"Story status updated to: {new_status}")
        return jsonify({
            'success': True,
            'message': 'Story status updated successfully',
            'story': story
        })
    except exceptions.CosmosResourceNotFoundError:
        print(f"Story not found with ID: {story_id}")
        return jsonify({'error': 'Story not found'}), 404
    except Exception as e:
        print(f"Error updating story status: {str(e)}")
        print("Stack trace:", traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@srs_brd_to_stories_bp.route('/edit-story', methods=['POST'])
def edit_story():
    try:
        print("\n=== Starting edit-story request ===")
        print("Request headers:", request.headers)
        print("Request data:", request.get_json())
        
        data = request.get_json()
        story_id = data.get('story_id')
        edit_type = data.get('edit_type')  # 'ai' or 'manual'
        story_data = data.get('story_data')  # For manual edits
        
        if not story_id:
            print("Error: Story ID is missing")
            return jsonify({'error': 'Story ID is required'}), 400
            
        if not story_data:
            print("Error: Story data is missing")
            return jsonify({'error': 'Story data is required'}), 400
            
        # Validate required fields
        required_fields = ['title', 'user_role', 'goal', 'benefit', 'acceptance_criteria']
        missing_fields = [field for field in required_fields if field not in story_data]
        if missing_fields:
            print(f"Error: Missing required fields: {missing_fields}")
            return jsonify({'error': f'Missing required fields: {", ".join(missing_fields)}'}), 400
            
        # Get database and container
        print("Connecting to Cosmos DB...")
        database = cosmos_client.get_database_client(COSMOS_DATABASE)
        user_stories_container = database.get_container_client("UserStories")
        
        # Get the existing story
        try:
            print(f"Fetching story with ID: {story_id}")
            # First try to find the story by querying
            query = f"SELECT * FROM c WHERE c.id = '{story_id}'"
            items = list(user_stories_container.query_items(
                query=query,
                enable_cross_partition_query=True
            ))
            
            if not items:
                print(f"Error: Story not found with ID: {story_id}")
                return jsonify({'error': 'Story not found'}), 404
                
            story = items[0]
            print("Story found:", story)
            
            if edit_type == 'manual':
                # Manual edit
                try:
                    print("Updating story with new data:", story_data)
                    # Update story with new data
                    for key, value in story_data.items():
                        if key in story:
                            story[key] = value
                    
                    # Update the story in the database
                    print("Saving updated story to Cosmos DB")
                    user_stories_container.replace_item(item=story_id, body=story)
                    print("Story updated successfully")
                    return jsonify({
                        'success': True,
                        'message': 'Story updated successfully',
                        'story': story
                    })
                except Exception as e:
                    print(f"Error updating story: {str(e)}")
                    print("Stack trace:", traceback.format_exc())
                    return jsonify({'error': str(e)}), 500
                    
            else:
                print(f"Error: Invalid edit type: {edit_type}")
                return jsonify({'error': 'Invalid edit type'}), 400
                
        except Exception as e:
            print(f"Error fetching story: {str(e)}")
            print("Stack trace:", traceback.format_exc())
            return jsonify({'error': str(e)}), 500
            
    except Exception as e:
        print(f"Error in edit-story route: {str(e)}")
        print("Stack trace:", traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@srs_brd_to_stories_bp.route('/push-to-jira/<story_id>', methods=['POST'])
def push_to_jira(story_id):
    try:
        print("\n=== Starting push-to-jira request ===")
        print(f"Story ID: {story_id}")
        
        # Get database and container
        print("Connecting to Cosmos DB...")
        database = cosmos_client.get_database_client(COSMOS_DATABASE)
        user_stories_container = database.get_container_client("UserStories")
        
        # Get the story
        try:
            print(f"Fetching story with ID: {story_id}")
            query = f"SELECT * FROM c WHERE c.id = '{story_id}'"
            items = list(user_stories_container.query_items(
                query=query,
                enable_cross_partition_query=True
            ))
            
            if not items:
                print(f"Error: Story not found with ID: {story_id}")
                return jsonify({'error': 'Story not found'}), 404
                
            story = items[0]
            print("Story found:", story)
            
            # Check if story already has a Jira issue ID
            if story.get('jira_issue_id'):
                print("Story already has a Jira issue ID:", story['jira_issue_id'])
                return jsonify({
                    'success': True,
                    'message': 'Story already pushed to Jira',
                    'jira_issue_id': story['jira_issue_id']
                })
            
            # Push to Jira
            try:
                print("Pushing story to Jira...")
                jira_issue_id = push_story_to_jira(story)
                
                if jira_issue_id:
                    # Update story with Jira issue ID
                    story['jira_issue_id'] = jira_issue_id
                    user_stories_container.replace_item(item=story_id, body=story)
                    
                    print(f"Successfully pushed to Jira with issue ID: {jira_issue_id}")
                    return jsonify({
                        'success': True,
                        'message': 'Story successfully pushed to Jira',
                        'jira_issue_id': jira_issue_id
                    })
                else:
                    print("Failed to push to Jira")
                    return jsonify({'error': 'Failed to push story to Jira'}), 500
                    
            except Exception as e:
                print(f"Error pushing to Jira: {str(e)}")
                print("Stack trace:", traceback.format_exc())
                return jsonify({'error': str(e)}), 500
                
        except Exception as e:
            print(f"Error fetching story: {str(e)}")
            print("Stack trace:", traceback.format_exc())
            return jsonify({'error': str(e)}), 500
            
    except Exception as e:
        print(f"Error in push-to-jira route: {str(e)}")
        print("Stack trace:", traceback.format_exc())
        return jsonify({'error': str(e)}), 500

if __name__ == "__main__":
    main()